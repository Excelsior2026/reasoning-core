"""Document parsers for various file formats."""

import os
from typing import Dict, Optional
from pathlib import Path


class DocumentParserError(Exception):
    """Exception raised when document parsing fails."""

    pass


def parse_document(file_path: str, file_type: Optional[str] = None) -> Dict[str, str]:
    """Parse a document and extract text content.

    Args:
        file_path: Path to the document file
        file_type: Optional file type hint (pdf, docx, txt, html, etc.)

    Returns:
        Dictionary with 'text' (extracted content) and 'metadata' (file info)

    Raises:
        DocumentParserError: If parsing fails
        FileNotFoundError: If file doesn't exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # Detect file type from extension if not provided
    if not file_type:
        ext = Path(file_path).suffix.lower()
        file_type = ext.lstrip(".")

    file_type = file_type.lower()

    try:
        if file_type == "pdf":
            return _parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return _parse_docx(file_path)
        elif file_type == "txt":
            return _parse_txt(file_path)
        elif file_type in ["html", "htm"]:
            return _parse_html(file_path)
        elif file_type in ["md", "markdown"]:
            return _parse_markdown(file_path)
        else:
            # Try to parse as plain text as fallback
            return _parse_txt(file_path)
    except Exception as e:
        raise DocumentParserError(f"Failed to parse {file_type} file: {e}") from e


def _parse_pdf(file_path: str) -> Dict[str, str]:
    """Parse PDF file."""
    try:
        import PyPDF2

        text_content = []
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())

        text = "\n".join(text_content)
        return {
            "text": text,
            "metadata": {
                "file_type": "pdf",
                "num_pages": num_pages,
                "file_size": os.path.getsize(file_path),
            },
        }
    except ImportError:
        raise DocumentParserError(
            "PyPDF2 not installed. Install with: pip install PyPDF2"
        )
    except Exception as e:
        raise DocumentParserError(f"PDF parsing failed: {e}") from e


def _parse_docx(file_path: str) -> Dict[str, str]:
    """Parse DOCX file."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_content = []

        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        text = "\n".join(text_content)
        return {
            "text": text,
            "metadata": {
                "file_type": "docx",
                "num_paragraphs": len(doc.paragraphs),
                "file_size": os.path.getsize(file_path),
            },
        }
    except ImportError:
        raise DocumentParserError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        raise DocumentParserError(f"DOCX parsing failed: {e}") from e


def _parse_txt(file_path: str) -> Dict[str, str]:
    """Parse plain text file."""
    try:
        # Try different encodings
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        text = None

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    text = file.read()
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise DocumentParserError("Could not decode file with any encoding")

        return {
            "text": text,
            "metadata": {
                "file_type": "txt",
                "file_size": os.path.getsize(file_path),
                "num_lines": len(text.splitlines()),
            },
        }
    except Exception as e:
        raise DocumentParserError(f"Text file parsing failed: {e}") from e


def _parse_html(file_path: str) -> Dict[str, str]:
    """Parse HTML file."""
    try:
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8") as file:
            html_content = file.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n", strip=True)

        # Extract title
        title = soup.find("title")
        title_text = title.string if title else ""

        return {
            "text": text,
            "metadata": {
                "file_type": "html",
                "title": title_text,
                "file_size": os.path.getsize(file_path),
            },
        }
    except ImportError:
        raise DocumentParserError(
            "beautifulsoup4 not installed. Install with: pip install beautifulsoup4"
        )
    except Exception as e:
        raise DocumentParserError(f"HTML parsing failed: {e}") from e


def _parse_markdown(file_path: str) -> Dict[str, str]:
    """Parse Markdown file."""
    # For now, treat markdown as plain text
    # Could use markdown library to convert to HTML then extract text
    return _parse_txt(file_path)
